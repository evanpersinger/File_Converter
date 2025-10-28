# docx_pdf.py
# converts docx files to pdf

import os
import sys
import argparse
from pathlib import Path
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document


# Create input and output directories if they don't exist
def setup_directories(input_dir="input", output_dir="output"):
    input_path = Path(input_dir)
    output_path = Path(output_dir)
    
    input_path.mkdir(exist_ok=True)
    output_path.mkdir(exist_ok=True)
    
    return input_path, output_path


# convert a docx table to a pdf table and add it to the story
def add_table_to_story(story, table):
    # Get table data
    data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            # Get text from cell, replacing newlines with spaces
            cell_text = cell.text.replace('\n', ' ')
            row_data.append(cell_text)
        data.append(row_data)
    
    # Create PDF table
    pdf_table = Table(data)
    
    # Style the table
    pdf_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('TOPPADDING', (0, 1), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    
    story.append(pdf_table)
    story.append(Spacer(1, 12))


# Convert a .docx file to PDF
# returns True if successful, False otherwise
def convert_docx_to_pdf(docx_path, output_path=None, input_dir=None, output_dir=None):
    # Set default directories if not provided
    if input_dir is None:
        input_dir = Path("input")
    else:
        input_dir = Path(input_dir)
    
    if output_dir is None:
        output_dir = Path("output")
    else:
        output_dir = Path(output_dir)

    # if the provided path already points to a PDF let user know and return False
    try:
        provided_suffix = Path(docx_path).suffix.lower()
        if provided_suffix == ".pdf":
            print("That file is already in pdf format")
            return False
    except Exception:
        pass

    # Build full input path
    full_input_path = Path(docx_path) if os.path.isabs(str(docx_path)) else input_dir / docx_path
    if not full_input_path.exists():
        print(f"Error: Word file '{full_input_path}' not found")
        return False

    # Compute output name
    if output_path is None:
        pdf_name = f"{full_input_path.stem}.pdf"
    else:
        pdf_name = Path(output_path).name
    full_output_path = output_dir / pdf_name

    try:
        # Load the Word document
        docx = Document(str(full_input_path))

        # Create PDF document
        doc = SimpleDocTemplate(str(full_output_path), pagesize=letter)
        styles = getSampleStyleSheet()

        # Create custom styles
        title_style = ParagraphStyle(
            'DocxTitle',
            parent=styles['Title'],
            fontSize=16,
            spaceAfter=20,
            alignment=1  # Center alignment
        )

        normal_style = ParagraphStyle(
            'DocxStyle',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,
            leftIndent=0,
            rightIndent=0,
            spaceAfter=6
        )

        # Build PDF content
        story = []

        # Add title
        filename = full_input_path.stem
        title = Paragraph(f"Document: {filename}", title_style)
        story.append(title)
        story.append(Spacer(1, 20))

        # Process document elements in order (paras and tables interleaved)
        # Create maps for quick lookup
        para_map = {id(para._p): para for para in docx.paragraphs}
        table_map = {id(table._tbl): table for table in docx.tables}
        
        # Iterate through body elements in document order
        for element in docx.element.body:
            elem_id = id(element)
            
            # Check if it's a paragraph
            if elem_id in para_map:
                para = para_map[elem_id]
                if para.text.strip():
                    # Escape special characters for reportlab
                    escaped_text = para.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(escaped_text, normal_style))
                else:
                    story.append(Spacer(1, 6))
            
            # Check if it's a table
            elif elem_id in table_map:
                story.append(Spacer(1, 12))
                add_table_to_story(story, table_map[elem_id])

        # Build PDF
        print(f"Converting '{full_input_path}' to '{full_output_path}'...")
        doc.build(story)
        
        if full_output_path.exists():
            print(f"Successfully converted to '{full_output_path}'")
            return True
        else:
            print("PDF creation failed")
            return False

    except Exception as e:
        print(f"Error creating PDF: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    parser = argparse.ArgumentParser(description="Convert DOCX files to PDF")
    parser.add_argument("docx_file", nargs="?", help="DOCX file to convert (optional)")
    parser.add_argument("output_file", nargs="?", help="Output PDF filename (optional)")
    parser.add_argument("--input-dir", default="input", help="Input directory (default: input)")
    parser.add_argument("--output-dir", default="output", help="Output directory (default: output)")
    
    args = parser.parse_args()
    
    # Set up directories
    input_dir, output_dir = setup_directories(args.input_dir, args.output_dir)
    
    if args.docx_file:
        # Convert specific file
        success = convert_docx_to_pdf(args.docx_file, args.output_file, input_dir, output_dir)
        if not success:
            sys.exit(1)
    else:
        # No args: convert all .docx files in input/
        docx_files = sorted(p for p in input_dir.glob("*.docx"))
        if not docx_files:
            print(f"No .docx files found in {input_dir} folder")
            print("Usage: python docx_pdf.py <file.docx> [output.pdf] [--input-dir DIR] [--output-dir DIR]")
            print("Example: python docx_pdf.py notes.docx")
            print("Example: python docx_pdf.py notes.docx my_notes.pdf")
            print("Example: python docx_pdf.py --input-dir myinput --output-dir myoutput")
            return
        any_failed = False
        for docx in docx_files:
            ok = convert_docx_to_pdf(docx.name, None, input_dir, output_dir)
            if not ok:
                any_failed = True
        if any_failed:
            sys.exit(1)


if __name__ == "__main__":
    main()
